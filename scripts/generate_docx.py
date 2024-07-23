import sys
from docx import Document
from docx.shared import Inches
import json

class DocxProcessor:
    def __init__(self, input_path, output_path, replacements, image_replacements):
        self.input_path = input_path
        self.output_path = output_path
        self.replacements = replacements
        self.image_replacements = image_replacements

    def replace_text(self):
        doc = Document(self.input_path)

        for paragraph in doc.paragraphs:
            for key, value in self.replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)
        for placeholder, image_path in self.image_replacements.items():
            self.replace_image(doc, placeholder, image_path)
        doc.save(self.output_path)

    def replace_image(self, doc, placeholder, image_path):
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                print(placeholder)
                parts = paragraph.text.split(placeholder)
                print(parts)
                if len(parts) > 1:
                    paragraph.clear()  
                    paragraph.add_run(parts[0]) 
                    run = paragraph.add_run()  
                    run.add_picture(image_path, width=Inches(0.25))  
                    paragraph.add_run(parts[1])  

if __name__ == "__main__":
    input_docx_path = sys.argv[1]
    output_docx_path = sys.argv[2]
    replacements = json.loads(sys.argv[3])
    image_replacements = json.loads(sys.argv[4])
    processor = DocxProcessor(input_docx_path, output_docx_path, replacements, image_replacements)
    processor.replace_text()
